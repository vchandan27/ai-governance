"""Extract plain text from uploaded policy documents.

Supported formats: .txt, .md, .pdf, .docx. The extractor is intentionally
defensive: unsupported or corrupt files raise ``ExtractionError`` with a
human-readable message rather than leaking stack traces to the API.
"""
from __future__ import annotations

import io
import re
from pathlib import Path


class ExtractionError(Exception):
    """Raised when a document cannot be parsed."""


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf_pymupdf(data: bytes) -> str | None:
    """Preferred extractor: PyMuPDF yields much cleaner text on justified PDFs
    (e.g. the EU Official Journal) where pypdf inserts spurious intra-word
    spaces. Returns None if PyMuPDF is unavailable so the caller can fall back."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return None
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc)
    except Exception:
        return None


def _extract_pdf_pypdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency guaranteed
        raise ExtractionError("PDF support is not installed.") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not read PDF: {exc}") from exc

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            # Skip pages that fail to extract rather than aborting the whole doc.
            continue
    return "\n".join(pages)


def _extract_pdf(data: bytes) -> str:
    text = _extract_pdf_pymupdf(data)
    if not text or not text.strip():
        text = _extract_pdf_pypdf(data)
    if not text.strip():
        raise ExtractionError(
            "No selectable text found in the PDF. It may be a scanned image; "
            "OCR is required (not enabled in this draft)."
        )
    return text


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency guaranteed
        raise ExtractionError("DOCX support is not installed.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not read DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" \t ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_html(data: bytes) -> str:
    raw = _extract_txt(data)
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        # Fallback: crude tag stripping if BeautifulSoup is unavailable.
        return re.sub(r"<[^>]+>", " ", raw)
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text("\n")


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch extraction based on the file extension."""
    suffix = Path(filename).suffix.lower()
    if suffix in (".txt", ".md"):
        text = _extract_txt(data)
    elif suffix == ".pdf":
        text = _extract_pdf(data)
    elif suffix == ".docx":
        text = _extract_docx(data)
    elif suffix in (".html", ".htm"):
        text = _extract_html(data)
    else:
        raise ExtractionError(f"Unsupported file type: '{suffix or filename}'.")

    text = normalize_whitespace(text)
    if not text.strip():
        raise ExtractionError("The document appears to be empty.")
    return text


def normalize_whitespace(text: str) -> str:
    # Collapse runs of spaces/tabs but preserve paragraph breaks.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def segment_text(text: str, min_len: int = 25) -> list[str]:
    """Split a policy document into analysable segments.

    We segment on paragraph boundaries and then further on sentence boundaries
    for long paragraphs, so that evidence snippets are concise and matching is
    granular. Very short fragments (headings, page numbers) are dropped.
    """
    segments: list[str] = []
    paragraphs = re.split(r"\n\s*\n", text)
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(para) <= 320:
            if len(para) >= min_len:
                segments.append(para)
            continue
        # Split long paragraphs into sentences.
        sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", para)
        buffer = ""
        for sent in sentences:
            sent = sent.strip()
            if not sent:
                continue
            if len(buffer) + len(sent) < 320:
                buffer = f"{buffer} {sent}".strip()
            else:
                if len(buffer) >= min_len:
                    segments.append(buffer)
                buffer = sent
        if len(buffer) >= min_len:
            segments.append(buffer)
    # De-duplicate while preserving order.
    seen: set[str] = set()
    unique: list[str] = []
    for seg in segments:
        key = seg.lower()
        if key not in seen:
            seen.add(key)
            unique.append(seg)
    return unique
